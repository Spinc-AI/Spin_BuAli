"""Spin BuAli — a Tkinter desktop client for the BuAli controller.

Single-window client (no tabs — this project does exactly one thing):
choose a pipeline (Separate / Multimodal / Hybrid), configure up to 3
independent STT slots (Separate/Hybrid) and/or an audio-capable LLM
(Multimodal/Hybrid), record or pick an audio file, and see the corrected
radiology report.

Ported from Spin_Medical_Assistant_Project/demo_app's shared widget patterns
(ModeSelector, CloudFieldsFrame, SttSlotWidget, ConnectionBar, MicRecorder/
RecordButton, TranscriptSaver, OutputBox, ScrollableFrame) and its
OrchestratorTab's pipeline/session/run flow, trimmed to just what BuAli needs.
"""
import json
import os
import queue
import re
import tempfile
import threading
import wave
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

try:
    import sounddevice as sd
    MIC_ERROR = None
except Exception as exc:  # missing PortAudio / no audio device — keep the app usable
    sd = None
    MIC_ERROR = exc

DEFAULT_HOST = "localhost"
DEFAULT_PORT = 9002
TIMEOUT_SHORT = 15
TIMEOUT_LOAD = 900   # model loading / inference can be slow
AUDIO_FILETYPES = [("Audio files", "*.wav *.mp3 *.flac *.ogg *.m4a"), ("All files", "*.*")]
MIC_SAMPLE_RATE = 16000

REMOTE_LOCAL_LABEL = "Remote Local Model"
CUSTOM_CLOUD_LABEL = "Custom Cloud API"
DEFAULT_CLOUD_BASE_URL = "https://api.openai.com/v1"
DEFAULT_CLOUD_STT_MODEL = "whisper-1"
DEFAULT_CLOUD_LLM_MODEL = "gpt-4o-mini"
MAX_STT_SLOTS = 3

LOCAL_LLM_MODELS = ["aya-expanse-8b", "aya-expanse-32b", "gemma-4-31b",
                    "gemma-4-e4b", "gemma-4-12b", "qwen3-omni-30b"]
LOCAL_AUDIO_MODELS = ["gemma-4-e4b", "gemma-4-12b", "qwen3-omni-30b"]

SEPARATE_LABEL = "Separate"
MULTIMODAL_LABEL = "Multimodal"
HYBRID_LABEL = "Hybrid (STT + Multimodal LLM)"
PIPELINE_LABEL_TO_VALUE = {SEPARATE_LABEL: "separate", MULTIMODAL_LABEL: "multimodal", HYBRID_LABEL: "hybrid"}


def error_detail(exc: requests.RequestException) -> str:
    if exc.response is not None:
        try:
            data = exc.response.json()
            return data.get("detail", str(data))
        except ValueError:
            return exc.response.text or str(exc)
    return str(exc)


def run_bg(fn, *args, **kwargs):
    threading.Thread(target=fn, args=args, kwargs=kwargs, daemon=True).start()


class ModeSelector(ttk.Frame):
    """Label + dropdown choosing "Remote Local Model" vs "Custom Cloud API"."""

    def __init__(self, parent, label_text):
        super().__init__(parent)
        ttk.Label(self, text=label_text).grid(row=0, column=0, sticky="w")
        self.mode = tk.StringVar(value=REMOTE_LOCAL_LABEL)
        self.box = ttk.Combobox(self, textvariable=self.mode, width=20, state="readonly",
                                values=[REMOTE_LOCAL_LABEL, CUSTOM_CLOUD_LABEL])
        self.box.grid(row=0, column=1, sticky="w", padx=(6, 0))

    def is_cloud(self):
        return self.mode.get() == CUSTOM_CLOUD_LABEL


class CloudFieldsFrame(ttk.Frame):
    """Cloud model / API key / base URL fields."""

    def __init__(self, parent, default_model):
        super().__init__(parent)
        ttk.Label(self, text="Cloud model:").grid(row=0, column=0, sticky="w")
        self.model = tk.StringVar(value=default_model)
        ttk.Entry(self, textvariable=self.model, width=18).grid(row=0, column=1, sticky="w", padx=(0, 12))

        ttk.Label(self, text="API key:").grid(row=0, column=2, sticky="w")
        self.api_key = tk.StringVar()
        ttk.Entry(self, textvariable=self.api_key, show="*", width=24).grid(row=0, column=3, sticky="w")

        ttk.Label(self, text="Base URL:").grid(row=1, column=0, sticky="w", pady=(4, 0))
        self.base_url = tk.StringVar(value=DEFAULT_CLOUD_BASE_URL)
        ttk.Entry(self, textvariable=self.base_url, width=40).grid(
            row=1, column=1, columnspan=3, sticky="w", pady=(4, 0))


class SttSlotWidget(ttk.Frame):
    """One independently local-or-cloud STT engine slot: an explicit "Use this
    slot" checkbox, a mode selector, and local dropdown / cloud fields."""

    def __init__(self, parent, label_text, local_models_getter, on_refresh=None):
        super().__init__(parent)
        self._local_models_getter = local_models_getter
        self._on_refresh = on_refresh

        self.enabled = tk.BooleanVar(value=False)
        ttk.Checkbutton(self, text="Use this slot", variable=self.enabled,
                        command=self._update_visibility).grid(row=0, column=0, sticky="w")

        self.mode = ModeSelector(self, label_text)
        self.mode.grid(row=0, column=1, columnspan=3, sticky="w", padx=(10, 0))
        self.mode.mode.trace_add("write", lambda *a: self._update_visibility())

        self.local_model = tk.StringVar()
        self.local_box = ttk.Combobox(self, textvariable=self.local_model, width=18, state="readonly")
        self.refresh_btn = ttk.Button(self, text="Refresh", command=self._refresh_clicked)

        self.cloud = CloudFieldsFrame(self, DEFAULT_CLOUD_STT_MODEL)

        self._update_visibility()

    def refresh_local_models(self):
        models = self._local_models_getter()
        self.local_box["values"] = models
        if self.local_model.get() not in models and models:
            self.local_model.set(models[0])

    def _refresh_clicked(self):
        if self._on_refresh:
            self._on_refresh()

    def _update_visibility(self):
        if not self.enabled.get():
            self.local_box.grid_remove()
            self.refresh_btn.grid_remove()
            self.cloud.grid_remove()
            return
        if self.mode.is_cloud():
            self.local_box.grid_remove()
            self.refresh_btn.grid_remove()
            self.cloud.grid(row=1, column=0, columnspan=4, sticky="w", pady=(4, 0))
        else:
            self.cloud.grid_remove()
            self.local_box.grid(row=1, column=0, sticky="w", pady=(4, 0))
            self.refresh_btn.grid(row=1, column=1, sticky="w", padx=(4, 0), pady=(4, 0))

    def is_cloud(self):
        return self.mode.is_cloud()

    def effective_model(self):
        if self.is_cloud():
            return "openai:" + self.cloud.model.get().strip()
        return self.local_model.get().strip()

    def as_slot_config(self):
        if not self.enabled.get():
            return None
        cfg = {"model": self.effective_model()}
        if self.is_cloud():
            key = self.cloud.api_key.get().strip()
            base = self.cloud.base_url.get().strip()
            if key:
                cfg["api_key"] = key
            if base:
                cfg["base_url"] = base
        return cfg


class MicRecorder:
    """Records mono 16-bit PCM from the default microphone until stopped."""

    def __init__(self, samplerate=MIC_SAMPLE_RATE, channels=1):
        self.samplerate = samplerate
        self.channels = channels
        self._queue = queue.Queue()
        self._stream = None

    def _callback(self, indata, frames, time_info, status):
        self._queue.put(bytes(indata))

    def start(self):
        self._queue = queue.Queue()
        self._stream = sd.RawInputStream(
            samplerate=self.samplerate, channels=self.channels,
            dtype="int16", callback=self._callback,
        )
        self._stream.start()

    def stop_and_save(self):
        self._stream.stop()
        self._stream.close()
        chunks = []
        while not self._queue.empty():
            chunks.append(self._queue.get())
        path = tempfile.NamedTemporaryFile(suffix=".wav", delete=False).name
        with wave.open(path, "wb") as wf:
            wf.setnchannels(self.channels)
            wf.setsampwidth(2)
            wf.setframerate(self.samplerate)
            wf.writeframes(b"".join(chunks))
        return path


class RecordButton(ttk.Button):
    def __init__(self, parent, file_path_var):
        super().__init__(parent, text="Record mic", command=self._toggle)
        self.file_path_var = file_path_var
        self._recorder = None

    def _toggle(self):
        if self._recorder is None:
            self._start()
        else:
            self._stop()

    def _start(self):
        if sd is None:
            messagebox.showerror("Microphone", f"Microphone support unavailable: {MIC_ERROR}")
            return
        try:
            self._recorder = MicRecorder()
            self._recorder.start()
        except Exception as exc:
            messagebox.showerror("Microphone", f"Could not start recording: {exc}")
            self._recorder = None
            return
        self.config(text="Stop recording")

    def _stop(self):
        try:
            path = self._recorder.stop_and_save()
            self.file_path_var.set(path)
        except Exception as exc:
            messagebox.showerror("Microphone", f"Could not save recording: {exc}")
        finally:
            self._recorder = None
            self.config(text="Record mic")


class ConnectionBar(ttk.Frame):
    def __init__(self, parent, default_port, health_path="/"):
        super().__init__(parent)
        self.health_path = health_path
        ttk.Label(self, text="Host:").grid(row=0, column=0, padx=(0, 4))
        self.host = tk.StringVar(value=DEFAULT_HOST)
        ttk.Entry(self, textvariable=self.host, width=16).grid(row=0, column=1, padx=(0, 8))
        ttk.Label(self, text="Port:").grid(row=0, column=2, padx=(0, 4))
        self.port = tk.StringVar(value=str(default_port))
        ttk.Entry(self, textvariable=self.port, width=6).grid(row=0, column=3, padx=(0, 8))
        ttk.Button(self, text="Check connection", command=self.check).grid(row=0, column=4, padx=(0, 8))
        self.indicator = ttk.Label(self, text="● unknown", foreground="gray")
        self.indicator.grid(row=0, column=5)

    @property
    def base_url(self):
        return f"http://{self.host.get().strip()}:{self.port.get().strip()}"

    def check(self):
        self._set("● checking...", "gray")
        run_bg(self._check_bg)

    def _check_bg(self):
        try:
            ok = requests.get(self.base_url + self.health_path, timeout=TIMEOUT_SHORT).status_code == 200
        except requests.RequestException:
            ok = False
        self.after(0, self._set, ("● connected" if ok else "● unreachable"), ("green" if ok else "red"))

    def _set(self, text, color):
        self.indicator.config(text=text, foreground=color)


# ---------------------------------------------------------------------------
# Transcript -> Word (.docx) export
# ---------------------------------------------------------------------------
# Persian/Arabic-block codepoint ranges. A line containing any of these is
# treated as an RTL paragraph.
_RTL_RANGES = [(0x0590, 0x08FF), (0xFB1D, 0xFDFF), (0xFE70, 0xFEFF)]
_FILENAME_UNSAFE_RE = re.compile(r'[\\/:*?"<>|]')


def _looks_rtl(text):
    return any(any(lo <= ord(ch) <= hi for lo, hi in _RTL_RANGES) for ch in (text or ""))


def sanitize_filename_part(name):
    if not name:
        return ""
    name = name[len("openai:"):] if name.startswith("openai:") else name
    name = _FILENAME_UNSAFE_RE.sub("-", name)
    return name.strip("-_")


def build_transcript_filename(audio_path, model_parts, ext=".docx"):
    stem = os.path.splitext(os.path.basename(audio_path))[0] if audio_path else "report"
    parts = [stem] + [sanitize_filename_part(m) for m in model_parts if m]
    return "_".join(p for p in parts if p) + ext


def save_text_as_docx(text, path):
    doc = Document()
    for line in (text or "").split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        if _looks_rtl(line):
            p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))
            run._r.get_or_add_rPr().append(OxmlElement("w:rtl"))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(path)


class TranscriptSaver(ttk.Frame):
    def __init__(self, parent, get_text_and_parts):
        super().__init__(parent)
        self._get_text_and_parts = get_text_and_parts
        self._overridden = False
        self.save_dir = tk.StringVar()

        ttk.Label(self, text="Save to:").grid(row=0, column=0, sticky="w")
        ttk.Entry(self, textvariable=self.save_dir, width=40, state="readonly").grid(
            row=0, column=1, sticky="w")
        ttk.Button(self, text="Browse...", command=self._browse).grid(row=0, column=2, sticky="w", padx=4)
        ttk.Button(self, text="Save Transcript (.docx)", command=self._save).grid(
            row=0, column=3, sticky="w", padx=(8, 0))

    def note_audio_path(self, audio_path):
        if audio_path and not self._overridden:
            self.save_dir.set(os.path.dirname(os.path.abspath(audio_path)))

    def _browse(self):
        d = filedialog.askdirectory(title="Choose where to save the transcript")
        if d:
            self.save_dir.set(d)
            self._overridden = True

    def _save(self):
        text, audio_path, model_parts = self._get_text_and_parts()
        if not text or not text.strip():
            messagebox.showwarning("Save Transcript", "Nothing to save yet — run BuAli first.")
            return
        save_dir = self.save_dir.get().strip() or (
            os.path.dirname(os.path.abspath(audio_path)) if audio_path else os.getcwd()
        )
        filename = build_transcript_filename(audio_path, model_parts)
        path = os.path.join(save_dir, filename)
        try:
            os.makedirs(save_dir, exist_ok=True)
            save_text_as_docx(text, path)
            messagebox.showinfo("Save Transcript", f"Saved to:\n{path}")
        except OSError as exc:
            messagebox.showerror("Save Transcript", f"Could not save: {exc}")


class OutputBox(scrolledtext.ScrolledText):
    def __init__(self, parent, height=12):
        super().__init__(parent, width=90, height=height, state="disabled", wrap="word")

    def write(self, text):
        self.config(state="normal")
        self.delete("1.0", tk.END)
        self.insert(tk.END, text)
        self.config(state="disabled")


class ScrollableFrame(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        canvas = tk.Canvas(self, highlightthickness=0)
        vsb = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")

        self.body = ttk.Frame(canvas, padding=10)
        window = canvas.create_window((0, 0), window=self.body, anchor="nw")

        self.body.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(window, width=e.width))

        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _on_mousewheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Spin BuAli — Demo")
        self.geometry("880x760")
        self.minsize(700, 500)

        outer = ScrollableFrame(self)
        outer.pack(fill="both", expand=True)
        root = outer.body

        self._last_audio_path = None
        self._last_result = {}

        self.conn = ConnectionBar(root, default_port=DEFAULT_PORT, health_path="/")
        self.conn.grid(row=0, column=0, columnspan=4, sticky="w", pady=(0, 10))

        # --- Pipeline ---
        ttk.Label(root, text="Pipeline:").grid(row=1, column=0, sticky="w")
        self.pipeline_label = tk.StringVar(value=SEPARATE_LABEL)
        self.pipeline_box = ttk.Combobox(root, textvariable=self.pipeline_label, width=28, state="readonly",
                                         values=[SEPARATE_LABEL, MULTIMODAL_LABEL, HYBRID_LABEL])
        self.pipeline_box.grid(row=1, column=1, sticky="w", padx=(6, 0))
        self.pipeline_label.trace_add("write", lambda *a: self._update_pipeline_visibility())

        # --- STT slots (separate/hybrid) ---
        self.slots_frame = ttk.LabelFrame(root, text="STT slots", padding=8)
        self.slots_frame.grid(row=2, column=0, columnspan=4, sticky="we", pady=(10, 0))
        self.slot_widgets = []
        for i in range(MAX_STT_SLOTS):
            w = SttSlotWidget(self.slots_frame, f"STT slot {i + 1} source:", self._get_local_stt_models,
                              on_refresh=self.refresh_stt_models)
            w.grid(row=i, column=0, sticky="w", pady=(4 if i else 0, 0))
            self.slot_widgets.append(w)

        # --- LLM ---
        self.llm_frame = ttk.LabelFrame(root, text="LLM", padding=8)
        self.llm_frame.grid(row=3, column=0, columnspan=4, sticky="we", pady=(10, 0))
        self.llm_mode = ModeSelector(self.llm_frame, "LLM source:")
        self.llm_mode.grid(row=0, column=0, columnspan=4, sticky="w")
        self.llm_mode.mode.trace_add("write", lambda *a: self._update_llm_section())

        self.llm_local_model = tk.StringVar(value=LOCAL_LLM_MODELS[0])
        self.llm_local_box = ttk.Combobox(self.llm_frame, textvariable=self.llm_local_model, width=18,
                                          state="readonly", values=LOCAL_LLM_MODELS)
        self.llm_cloud = CloudFieldsFrame(self.llm_frame, DEFAULT_CLOUD_LLM_MODEL)
        self.gemini_hint = ttk.Label(
            self.llm_frame,
            text="Multimodal/Hybrid: prefix the cloud model with \"gemini:\" to use Gemini's native audio API "
                 "(e.g. gemini:gemini-2.0-flash-exp) instead of the OpenAI-compatible shape.",
            foreground="gray", wraplength=520, justify="left",
        )

        # --- session controls ---
        ttk.Button(root, text="Start session", command=self.start_session).grid(
            row=4, column=0, sticky="w", pady=10)
        ttk.Button(root, text="Unload session", command=self.unload_session).grid(row=4, column=1, sticky="w")
        self.session_status = ttk.Label(root, text="session: none")
        self.session_status.grid(row=5, column=0, columnspan=4, sticky="w")

        # --- audio input ---
        ttk.Label(root, text="Audio file:").grid(row=6, column=0, sticky="w", pady=(10, 0))
        self.file_path = tk.StringVar()
        ttk.Entry(root, textvariable=self.file_path, width=45, state="readonly").grid(
            row=6, column=1, columnspan=2, sticky="w", pady=(10, 0))
        ttk.Button(root, text="Browse...", command=self.browse).grid(row=6, column=3, sticky="w", pady=(10, 0))
        RecordButton(root, self.file_path).grid(row=7, column=0, sticky="w", pady=(4, 0))
        ttk.Button(root, text="Clear", command=lambda: self.file_path.set("")).grid(
            row=7, column=1, sticky="w", pady=(4, 0))

        ttk.Button(root, text="Run", command=self.run).grid(row=8, column=0, sticky="w", pady=10)

        ttk.Label(root, text="Output:").grid(row=9, column=0, sticky="nw")
        self.output = OutputBox(root)
        self.output.grid(row=10, column=0, columnspan=4, pady=(0, 10))

        self.saver = TranscriptSaver(root, self._get_text_and_parts)
        self.saver.grid(row=11, column=0, columnspan=4, sticky="w")

        self._update_pipeline_visibility()
        self._update_llm_section()
        self.refresh_stt_models()

    # --- pipeline value ---
    def _pipeline_value(self):
        return PIPELINE_LABEL_TO_VALUE[self.pipeline_label.get()]

    def _update_pipeline_visibility(self):
        pipeline = self._pipeline_value()
        if pipeline == "multimodal":
            self.slots_frame.grid_remove()
        else:
            self.slots_frame.grid()
        self._update_llm_section()

    # --- LLM section ---
    def _update_llm_section(self):
        needs_audio = self._pipeline_value() in ("multimodal", "hybrid")
        values = LOCAL_AUDIO_MODELS if needs_audio else LOCAL_LLM_MODELS
        self.llm_local_box["values"] = values
        if self.llm_local_model.get() not in values:
            self.llm_local_model.set(values[0])

        if self.llm_mode.is_cloud():
            self.llm_local_box.grid_remove()
            self.llm_cloud.grid(row=1, column=0, columnspan=4, sticky="w", pady=(4, 0))
        else:
            self.llm_cloud.grid_remove()
            self.llm_local_box.grid(row=1, column=0, sticky="w", pady=(4, 0))

        if needs_audio and self.llm_mode.is_cloud():
            self.gemini_hint.grid(row=2, column=0, columnspan=4, sticky="w", pady=(4, 0))
        else:
            self.gemini_hint.grid_remove()

    def _get_local_stt_models(self):
        return self._local_stt_models if hasattr(self, "_local_stt_models") else []

    def refresh_stt_models(self):
        run_bg(self._refresh_stt_models_bg)

    def _refresh_stt_models_bg(self):
        try:
            r = requests.get(f"{self.conn.base_url}/models", timeout=TIMEOUT_SHORT)
            r.raise_for_status()
            self._local_stt_models = r.json().get("available", [])
            self.after(0, self._apply_stt_models)
        except requests.RequestException as exc:
            self.after(0, messagebox.showerror, "BuAli",
                      f"Could not fetch local STT models: {error_detail(exc)}")

    def _apply_stt_models(self):
        for w in self.slot_widgets:
            w.refresh_local_models()

    # --- effective LLM model ---
    def _effective_llm_model(self):
        if self.llm_mode.is_cloud():
            model = self.llm_cloud.model.get().strip()
            if model.startswith("gemini:"):
                return model
            return "openai:" + model
        return self.llm_local_model.get().strip()

    # --- session ---
    def start_session(self):
        payload = {"pipeline": self._pipeline_value(), "llm_model": self._effective_llm_model()}
        if self._pipeline_value() != "multimodal":
            slots = [w.as_slot_config() for w in self.slot_widgets]
            payload["stt_slots"] = slots
        if self.llm_mode.is_cloud():
            key = self.llm_cloud.api_key.get().strip()
            base = self.llm_cloud.base_url.get().strip()
            if key:
                payload["llm_api_key"] = key
            if base:
                payload["llm_base_url"] = base
        self.session_status.config(text="starting session...")
        run_bg(self._start_session_bg, payload)

    def _start_session_bg(self, payload):
        try:
            r = requests.post(f"{self.conn.base_url}/session", json=payload, timeout=TIMEOUT_LOAD)
            r.raise_for_status()
            data = r.json()
            text = f"session: pipeline={data.get('pipeline')}, llm={data.get('llm_model')}"
            self.after(0, self.session_status.config, {"text": text})
        except requests.RequestException as exc:
            self.after(0, self.session_status.config, {"text": f"session failed: {error_detail(exc)}"})

    def unload_session(self):
        run_bg(self._unload_session_bg)

    def _unload_session_bg(self):
        try:
            r = requests.post(f"{self.conn.base_url}/session/unload", timeout=TIMEOUT_SHORT)
            r.raise_for_status()
            self.after(0, self.session_status.config, {"text": "session: none"})
        except requests.RequestException as exc:
            self.after(0, messagebox.showerror, "BuAli", f"Unload failed: {error_detail(exc)}")

    # --- run ---
    def browse(self):
        path = filedialog.askopenfilename(title="Choose an audio file", filetypes=AUDIO_FILETYPES)
        if path:
            self.file_path.set(path)

    def run(self):
        path = self.file_path.get()
        if not path:
            messagebox.showwarning("BuAli", "Choose or record an audio file first.")
            return
        self._last_audio_path = path
        self.saver.note_audio_path(path)

        llm_key = self.llm_cloud.api_key.get().strip() if self.llm_mode.is_cloud() else ""
        llm_base = self.llm_cloud.base_url.get().strip() if self.llm_mode.is_cloud() else ""
        self.output.write("running...")
        run_bg(self._run_bg, path, llm_key, llm_base)

    def _run_bg(self, path, llm_key, llm_base):
        try:
            with open(path, "rb") as f:
                files = {"file": (os.path.basename(path), f)}
                data = {}
                if llm_key:
                    data["llm_api_key"] = llm_key
                if llm_base:
                    data["llm_base_url"] = llm_base
                r = requests.post(f"{self.conn.base_url}/run", files=files, data=data or None,
                                  timeout=TIMEOUT_LOAD)
            r.raise_for_status()
            body = r.json()
            self._last_result = body.get("result", {})
            pretty = json.dumps(body, indent=2, ensure_ascii=False)
            self.after(0, self.output.write, pretty)
        except requests.RequestException as exc:
            self.after(0, self.output.write, f"Error: {error_detail(exc)}")

    def _get_text_and_parts(self):
        text = self._last_result.get("final_text") or self._last_result.get("corrected_transcript") or ""
        model_parts = [self._effective_llm_model()]
        return text, self._last_audio_path, model_parts


if __name__ == "__main__":
    App().mainloop()
